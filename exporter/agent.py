from __future__ import annotations

import json
import logging
from pathlib import Path
from datetime import datetime

from docx import Document

log = logging.getLogger(__name__)


class Agent:
    """Exporter Agent capable of writing chat canvases to ``.docx`` files."""

    def run(
        self,
        canvas_json_path: str,
        template: str = "desktop_app/resources/CodexMechanica.dotx",
        out_dir: str | None = None,
    ) -> str:
        """Convert ``canvas_json_path`` to a Word document.

        Parameters
        ----------
        canvas_json_path:
            Path to the exported canvas JSON.
        template:
            Optional Word template to start from. Defaults to ``CodexMechanica.dotx``.
        out_dir:
            Directory where the ``.docx`` should be saved. Defaults to ``exports``.

        Returns
        -------
        str
            Absolute path to the generated document.
        """

        log.setLevel(logging.INFO)

        p = Path(canvas_json_path).expanduser().resolve()
        if not p.exists():
            raise FileNotFoundError(p)
        data = json.loads(p.read_text(encoding="utf-8"))

        canvas_id = data.get("id", f"canvas_{p.stem}_{datetime.utcnow().isoformat()}")
        out_root = Path(out_dir or "exports").resolve()
        out_root.mkdir(parents=True, exist_ok=True)

        template_path = Path(template)
        if template_path.exists():
            doc = Document(str(template_path))
        else:
            doc = Document()

        for block in data.get("blocks", []):
            author = block.get("author", "Unknown")
            text = block.get("text", "")
            doc.add_paragraph(f"{author}: {text}")

        out_path = out_root / f"{canvas_id}.docx"
        doc.save(out_path)
        log.info("Exported %s → %s", canvas_json_path, out_path)
        return str(out_path)

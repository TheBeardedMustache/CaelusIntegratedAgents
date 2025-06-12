"""Exporter agent capable of converting canvas JSON to Word docs."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from docx import Document


class Agent:
    """Simple exporter converting chat canvases into ``.docx`` files."""

    def run(self, canvas_json_path: str, template: str = "CodexMechanica.dotx") -> str:
        """Convert ``canvas_json_path`` to a Word document.

        Parameters
        ----------
        canvas_json_path:
            Path to a JSON file describing the canvas.
        template:
            Optional Word template to load. Defaults to ``"CodexMechanica.dotx"``.

        Returns
        -------
        str
            Absolute path to the generated ``.docx`` file.
        """

        try:
            logging.info("Loading canvas from %s", canvas_json_path)
            with open(canvas_json_path, "r", encoding="utf-8") as f:
                canvas: dict[str, Any] = json.load(f)
        except Exception as exc:  # pragma: no cover - exceptional path
            logging.error("Failed to load canvas JSON: %s", exc)
            raise

        canvas_id = str(canvas.get("id") or Path(canvas_json_path).stem)

        try:
            logging.info("Creating document using template %s", template)
            doc = Document(template)

            for block in canvas.get("blocks", []):
                if block.get("type") != "chat":
                    continue
                author = block.get("author", "")
                text = block.get("text", "")
                paragraph = doc.add_paragraph()
                if author:
                    paragraph.add_run(f"{author}: ").bold = True
                paragraph.add_run(text)

            export_dir = Path("exports")
            export_dir.mkdir(exist_ok=True)
            output_path = export_dir / f"{canvas_id}.docx"
            doc.save(output_path)
            abs_path = str(output_path.resolve())
            logging.info("Saved document to %s", abs_path)
            return abs_path
        except Exception as exc:  # pragma: no cover - exceptional path
            logging.error("Failed to export canvas %s: %s", canvas_id, exc)
            raise

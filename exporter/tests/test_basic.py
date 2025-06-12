import json
from pathlib import Path

from docx import Document

from exporter import Agent


def test_run_exports_docx(tmp_path, monkeypatch):
    """Agent should convert chat blocks into a Word document."""
    canvas = {
        "id": "42",
        "blocks": [
            {"type": "chat", "author": "user", "text": "hi"},
            {"type": "chat", "author": "assistant", "text": "hello"},
        ],
    }

    canvas_path = tmp_path / "canvas.json"
    canvas_path.write_text(json.dumps(canvas))

    template_path = tmp_path / "blank.dotx"
    Document().save(template_path)

    monkeypatch.chdir(tmp_path)
    agent = Agent()
    output = agent.run(str(canvas_path), template=str(template_path))

    expected = tmp_path / "exports" / "42.docx"
    assert Path(output) == expected
    assert expected.exists()

    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["user: hi", "assistant: hello"]
